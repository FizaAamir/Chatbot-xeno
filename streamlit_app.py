import streamlit as st
import fitz  # PyMuPDF
from docx import Document as DocxDocument # python-docx
from langchain.chains import create_retrieval_chain, create_history_aware_retriever
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import AIMessage, HumanMessage
from langchain_core.documents import Document
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
import os

# Initialize language model
os.environ['GOOGLE_API_KEY'] = 'Gemini-api-key'
llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.7, top_p=0.85)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

# System prompt
system_prompt = (
    "You are Xeno, an AI assistant that answers questions based on uploaded files. "
    "The user has uploaded the following files: {file_list}. "
    "If the user references 'both files', 'all documents', or 'uploaded PDFs', replace it with the actual file names from the list. "
    "Use the retrieved context to provide accurate responses. If the question requires multiple documents, combine relevant information. "
    "Always return the file name or path where the information was found. Keep responses concise."
    "\n\n{context}"
)

# Prompt templates
contextualize_q_prompt = ChatPromptTemplate.from_messages([
    ("system", 
        "You are an AI assistant that helps clarify user queries. "
        "The user uploaded the following files: {file_list}. "
        "If they say 'both files', 'all documents', or similar terms, replace it with the actual file names. "
        "Do NOT answer the question—just rewrite it to be more explicit."
        "Rephrase the user question so it is standalone, without answering it."
    ),
    MessagesPlaceholder("chat_history"),
    ("human", "{input}"),
])

qa_prompt = ChatPromptTemplate.from_messages([
    ("system", system_prompt),
    MessagesPlaceholder("chat_history"),
    ("human", "{input}"),
])

st.set_page_config(page_title="Xeno - Your AI File Assistant", page_icon=":robot:")
st.header("Xeno - Chat with Your Documents :books:")

# File upload section
st.sidebar.header("📂 Upload Documents")
uploaded_files = st.sidebar.file_uploader(
    "Upload PDFs or Word Docs", 
    type=["pdf", "docx"], 
    accept_multiple_files=True
)

# Store vectorized documents
documents = []
file_sources = {}  # To store file paths for responses

def extract_text_from_pdf(file):
    """Extracts text from a PDF file."""
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as pdf_doc:
        for page_num in range(pdf_doc.page_count):
            page = pdf_doc[page_num]
            text += page.get_text()
    return text

def extract_text_from_docx(file):
    """Extracts text from a Word document."""
    text = ""
    doc = DocxDocument(file)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

if uploaded_files:
    for uploaded_file in uploaded_files:
        text = ""
        if uploaded_file.type == "application/pdf":
            text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_from_docx(uploaded_file)

        if text:
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
            splits = text_splitter.split_text(text)

            for split in splits:
                doc = Document(page_content=split, metadata={"source": uploaded_file.name})
                documents.append(doc)
                file_sources[split] = uploaded_file.name  # Map chunks to file names


def determine_k(query, num_files):
    """Use LLM to determine optimal k for retrieval."""
    system_message = (
        "You are an AI assistant optimizing retrieval for a RAG-based chatbot. "
        "The chatbot retrieves relevant text chunks from uploaded PDF files to answer user queries. "
        "Based on the user's query and the total number of uploaded files, return an optimal 'k' value "
        "for retrieving document chunks. A higher 'k' is needed for broader, summarization-based queries, "
        "while a lower 'k' is sufficient for specific file references.\n\n"
        "Rules:\n"
        "- If the query is broad set k high (e.g., 5-7 chunks per file).\n"
        "- If the query is about a specific document, set k lower (e.g., 3-5 chunks).\n"
        "- Always return an integer k value without explanation.\n\n"
        "Example Inputs & Outputs:\n"
        "User query: 'Summarize all uploaded PDFs'\nTotal files: 10\nOutput: 50\n\n"
        "User query: 'Tell me about file X.pdf'\nTotal files: 10\nOutput: 5\n\n"
        "User query: 'What are the main differences between the files?'\nTotal files: 3\nOutput: 15\n\n"
    )

    response = llm.invoke(f"{system_message}\nUser query: '{query}'\nTotal files: {num_files}\nOutput:")
    
    # Extracting integer k from LLM response
    try:
        k_value = int(response.content.strip())
    except ValueError:
        k_value = max(num_files * 3, 15)  # Fallback to heuristic if LLM fails
    
    print(f"🔍 LLM-predicted k value: {k_value}")  # Debugging k value
    return k_value


if "retrieval_k" not in st.session_state:
    st.session_state.retrieval_k = 10

# Vector store for all documents
if documents:
    file_list = ", ".join([file.name for file in uploaded_files])  # Store uploaded file names
    vectorstore = InMemoryVectorStore.from_documents(documents=documents, embedding=embeddings)

    # Only initialize retriever without prompt-dependent k value
    retriever = vectorstore.as_retriever(search_kwargs={"k": st.session_state.retrieval_k})  

    history_aware_retriever = create_history_aware_retriever(
        llm, retriever, contextualize_q_prompt
    )
    question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)
    rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)
    
    st.session_state.rag_chain = rag_chain
    st.session_state.file_list = file_list  
    st.session_state.vectorstore = vectorstore  # Store vectorstore for dynamic retrieval update


# Initialize chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# Display chat history
for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Handle user query
if prompt := st.chat_input("Ask Xeno anything about your files..."):
    if "rag_chain" in st.session_state:
        # Determine dynamic retrieval k based on user query
        st.session_state.retrieval_k = max(determine_k(prompt, len(uploaded_files)), 10)
        retriever = st.session_state.vectorstore.as_retriever(search_kwargs={"k": st.session_state.retrieval_k})

        # Check if user is asking for file names
        if "list uploaded files" in prompt.lower() or "what files are uploaded" in prompt.lower():
            response_text = f"Uploaded PDFs: {st.session_state.file_list}"
        else:
            rag_chain = st.session_state.rag_chain
            ai_response = rag_chain.invoke({
                "input": prompt,
                "chat_history": st.session_state.chat_history,
                "file_list": st.session_state.file_list,
            })
            response_text = ai_response["answer"]
            
            # Extract sources from retrieved docs
            retrieved_docs = ai_response.get("context", [])
            print(f"Retrieved {len(retrieved_docs)} docs for query: {prompt}")
            # for doc in retrieved_docs:
            #     print(f" - Retrieved from {doc.metadata['source']} : {doc.page_content[:200]}")
            sources = set(file_sources.get(doc.page_content, "Unknown") for doc in retrieved_docs)
            # if sources:
            #     response_text += "\n\n**Sources:** " + ", ".join(sources)

        # Update chat history
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        st.session_state.chat_history.append({"role": "assistant", "content": response_text})

        # Display response
        with st.chat_message("user"):
            st.markdown(prompt)
        with st.chat_message("assistant"):
            st.markdown(response_text)
    else:
        st.warning("Please upload PDF files first.")
